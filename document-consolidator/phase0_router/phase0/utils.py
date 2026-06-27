"""
phase0/utils.py
Shared utilities: file text extraction, page counting, token estimation.
Supports .pdf, .docx, .txt, .xlsx (text content only for xlsx).
"""

from __future__ import annotations
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Rough chars-per-token estimate for context window budgeting
CHARS_PER_TOKEN = 4


def extract_text_from_file(filepath: str | Path) -> str:
    """
    Extracts plain text from a file based on its extension.
    Returns empty string on failure (logged as error).
    """
    filepath = Path(filepath)
    ext = filepath.suffix.lower()

    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".txt":  _extract_txt,
        ".md":   _extract_txt,
        ".xlsx": _extract_xlsx,
        ".csv":  _extract_txt,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        logger.warning(f"Unsupported file type: {ext} for {filepath.name}")
        return ""

    try:
        return extractor(filepath)
    except Exception as e:
        logger.error(f"Text extraction failed for {filepath.name}: {e}")
        return ""


def count_pages(filepath: str | Path) -> int:
    """Returns approximate page count. Returns 0 for non-paginated formats."""
    filepath = Path(filepath)
    ext = filepath.suffix.lower()

    if ext == ".pdf":
        try:
            import pypdf
            with open(filepath, "rb") as f:
                reader = pypdf.PdfReader(f)
                return len(reader.pages)
        except Exception:
            return 0

    if ext == ".docx":
        # Approximate: 500 words per page
        text = _extract_docx(filepath)
        word_count = len(text.split())
        return max(1, word_count // 500)

    return 0


def estimate_tokens(text: str) -> int:
    """Rough token count estimate — used for context window budgeting."""
    return len(text) // CHARS_PER_TOKEN


def page_ref_from_char_offset(char_offset: int, chars_per_page: int = 3000) -> str:
    """
    Converts a character offset in extracted text to an approximate page reference.
    Assumes ~3000 chars per page (typical dense document).
    """
    page = max(1, char_offset // chars_per_page + 1)
    return f"p.{page}"


# ---------------------------------------------------------------------------
# Private extractors
# ---------------------------------------------------------------------------

def _extract_pdf(filepath: Path) -> str:
    try:
        import pypdf
        text_parts = []
        with open(filepath, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        logger.error("pypdf not installed. Run: pip install pypdf")
        return ""


def _extract_docx(filepath: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""


def _extract_txt(filepath: Path) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            return filepath.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    logger.error(f"Could not decode {filepath.name} with any supported encoding")
    return ""


def _extract_xlsx(filepath: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        logger.error("openpyxl not installed. Run: pip install openpyxl")
        return ""
